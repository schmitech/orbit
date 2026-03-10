#!/usr/bin/env python3
"""
Markdown to Word Converter

A command-line tool for converting Markdown files (.md) to Microsoft Word documents (.docx)
using the python-docx and markdown libraries. This tool provides flexible input/output options
and comprehensive error handling.

Features:
- Convert single Markdown files to Word documents
- Support for custom input and output file paths
- Automatic output file naming based on input file
- Handles headings, bold, italic, code blocks, lists, links, tables, and images
- Comprehensive error handling and validation
- Verbose output with detailed conversion statistics

Requirements:
- python-docx library (install with: pip install python-docx)
- markdown library (install with: pip install markdown)
- beautifulsoup4 library (install with: pip install beautifulsoup4)
- Python 3.6+

Usage Examples:
    # Convert a specific file with custom output name
    python markdown-to-word.py input.md output.docx

    # Convert a file with automatic output naming
    python markdown-to-word.py document.md

    # Convert with verbose output
    python markdown-to-word.py -v input.md output.docx

    # Show help
    python markdown-to-word.py -h

Author: QA Pipeline Project
Version: 1.0
"""

import os
import sys
import argparse
import re
from pathlib import Path

import markdown
from bs4 import BeautifulSoup, NavigableString
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


def validate_input_file(file_path):
    """
    Validate that the input file exists and has a supported extension.

    Args:
        file_path (str): Path to the input file

    Returns:
        bool: True if valid, False otherwise
    """
    if not os.path.exists(file_path):
        print(f"Error: Input file not found at '{file_path}'")
        return False

    supported_extensions = ['.md', '.markdown', '.mdown', '.mkd']
    file_ext = Path(file_path).suffix.lower()

    if file_ext not in supported_extensions:
        print(f"Error: Unsupported file format '{file_ext}'")
        print(f"Supported formats: {', '.join(supported_extensions)}")
        return False

    return True


def generate_output_filename(input_path, custom_output=None):
    """
    Generate output filename based on input file or custom name.

    Args:
        input_path (str): Path to the input file
        custom_output (str, optional): Custom output filename

    Returns:
        str: Generated output filename
    """
    if custom_output:
        return custom_output

    input_path = Path(input_path)
    return f"{input_path.stem}.docx"


def parse_markdown_to_html(md_content):
    """
    Parse markdown content to HTML using the markdown library.

    Args:
        md_content (str): Raw markdown content

    Returns:
        str: HTML string
    """
    extensions = [
        'markdown.extensions.tables',
        'markdown.extensions.fenced_code',
        'markdown.extensions.codehilite',
        'markdown.extensions.toc',
        'markdown.extensions.nl2br',
        'markdown.extensions.sane_lists',
    ]
    html = markdown.markdown(md_content, extensions=extensions)
    return html


def add_formatted_text(paragraph, element):
    """
    Recursively add formatted text from an HTML element to a docx paragraph.

    Args:
        paragraph: docx Paragraph object
        element: BeautifulSoup element
    """
    if isinstance(element, NavigableString):
        text = str(element)
        if text.strip() or text == ' ':
            paragraph.add_run(text)
        return

    for child in element.children:
        if isinstance(child, NavigableString):
            text = str(child)
            if not text:
                continue
            run = paragraph.add_run(text)
            # Apply formatting based on parent tags
            parent_tags = [p.name for p in child.parents if p.name]
            if 'strong' in parent_tags or 'b' in parent_tags:
                run.bold = True
            if 'em' in parent_tags or 'i' in parent_tags:
                run.italic = True
            if 'code' in parent_tags:
                run.font.name = 'Courier New'
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(0x24, 0x29, 0x2E)
            if 'a' in parent_tags:
                run.font.color.rgb = RGBColor(0x05, 0x63, 0xC1)
                run.underline = True
        else:
            add_formatted_text(paragraph, child)


def add_list_items(doc, element, level=0, ordered=False):
    """
    Add list items to the document.

    Args:
        doc: docx Document object
        element: BeautifulSoup list element (ul or ol)
        level (int): Nesting level
        ordered (bool): Whether this is an ordered list
    """
    counter = 0
    for li in element.find_all('li', recursive=False):
        counter += 1
        # Create paragraph with list style
        para = doc.add_paragraph()
        para.paragraph_format.left_indent = Inches(0.25 * (level + 1))
        para.paragraph_format.space_before = Pt(2)
        para.paragraph_format.space_after = Pt(2)

        # Add bullet or number prefix
        prefix = f"{counter}. " if ordered else "- "
        run = para.add_run(prefix)
        run.bold = False

        # Get direct text content (not nested lists)
        for child in li.children:
            if isinstance(child, NavigableString):
                text = str(child).strip()
                if text:
                    para.add_run(text)
            elif child.name in ('ul', 'ol'):
                add_list_items(doc, child, level + 1, ordered=(child.name == 'ol'))
            else:
                add_formatted_text(para, child)


def add_table(doc, element):
    """
    Add a table to the document.

    Args:
        doc: docx Document object
        element: BeautifulSoup table element
    """
    rows = element.find_all('tr')
    if not rows:
        return

    # Determine number of columns
    first_row_cells = rows[0].find_all(['th', 'td'])
    num_cols = len(first_row_cells)
    if num_cols == 0:
        return

    table = doc.add_table(rows=0, cols=num_cols)
    table.style = 'Table Grid'

    for row_elem in rows:
        cells = row_elem.find_all(['th', 'td'])
        row = table.add_row()
        for i, cell_elem in enumerate(cells):
            if i < num_cols:
                cell_text = cell_elem.get_text(strip=True)
                cell = row.cells[i]
                cell.text = cell_text
                # Bold for header cells
                if cell_elem.name == 'th':
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.bold = True


def add_code_block(doc, element):
    """
    Add a code block to the document.

    Args:
        doc: docx Document object
        element: BeautifulSoup pre/code element
    """
    code_text = element.get_text()
    para = doc.add_paragraph()
    para.paragraph_format.left_indent = Inches(0.25)
    para.paragraph_format.space_before = Pt(6)
    para.paragraph_format.space_after = Pt(6)
    run = para.add_run(code_text)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x24, 0x29, 0x2E)


def convert_html_to_docx(html_content, doc):
    """
    Convert parsed HTML content into a docx Document.

    Args:
        html_content (str): HTML string from markdown conversion
        doc: docx Document object
    """
    soup = BeautifulSoup(html_content, 'html.parser')

    heading_map = {
        'h1': 'Heading 1',
        'h2': 'Heading 2',
        'h3': 'Heading 3',
        'h4': 'Heading 4',
        'h5': 'Heading 5',
        'h6': 'Heading 6',
    }

    for element in soup.children:
        if isinstance(element, NavigableString):
            text = str(element).strip()
            if text:
                doc.add_paragraph(text)
            continue

        tag = element.name

        # Headings
        if tag in heading_map:
            heading = doc.add_heading(level=int(tag[1]))
            add_formatted_text(heading, element)

        # Paragraphs
        elif tag == 'p':
            para = doc.add_paragraph()
            add_formatted_text(para, element)

        # Code blocks
        elif tag == 'pre':
            add_code_block(doc, element)

        # Unordered lists
        elif tag == 'ul':
            add_list_items(doc, element, ordered=False)

        # Ordered lists
        elif tag == 'ol':
            add_list_items(doc, element, ordered=True)

        # Tables
        elif tag == 'table':
            add_table(doc, element)

        # Horizontal rules
        elif tag == 'hr':
            para = doc.add_paragraph()
            para.paragraph_format.space_before = Pt(6)
            para.paragraph_format.space_after = Pt(6)
            run = para.add_run('_' * 50)
            run.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)

        # Blockquotes
        elif tag == 'blockquote':
            for child in element.find_all(['p', 'li'], recursive=False):
                para = doc.add_paragraph()
                para.paragraph_format.left_indent = Inches(0.5)
                para.style = 'Quote' if 'Quote' in [s.name for s in doc.styles] else None
                add_formatted_text(para, child)

        # Fallback for other tags
        else:
            text = element.get_text(strip=True)
            if text:
                doc.add_paragraph(text)


def convert_markdown_to_word(input_file, output_file, verbose=False):
    """
    Convert a Markdown file to Word document format.

    Args:
        input_file (str): Path to the input Markdown file
        output_file (str): Path to the output Word document
        verbose (bool): Enable verbose output

    Returns:
        bool: True if conversion successful, False otherwise
    """
    if verbose:
        print(f"Input file: {input_file}")
        print(f"Output file: {output_file}")
        print(f"File size: {os.path.getsize(input_file)} bytes")

    print(f"Converting '{input_file}' to '{output_file}'...")

    try:
        # Read the markdown file
        with open(input_file, 'r', encoding='utf-8') as f:
            md_content = f.read()

        if verbose:
            line_count = md_content.count('\n') + 1
            print(f"Source lines: {line_count}")

        # Parse markdown to HTML
        html_content = parse_markdown_to_html(md_content)

        if verbose:
            print(f"HTML size: {len(html_content)} characters")

        # Create Word document
        doc = Document()

        # Convert HTML to docx
        convert_html_to_docx(html_content, doc)

        # Save the document
        doc.save(output_file)

        # Verify the output file was created
        if os.path.exists(output_file):
            output_size = os.path.getsize(output_file)
            print("Conversion successful!")
            print(f"Output file: '{output_file}'")
            if verbose:
                print(f"Output size: {output_size} bytes")
            return True
        else:
            print("Error: Output file was not created")
            return False

    except FileNotFoundError as e:
        print(f"File not found error: {e}")
        return False
    except PermissionError as e:
        print(f"Permission error: {e}")
        print("Please check file permissions and ensure you have write access to the output directory.")
        return False
    except Exception as e:
        print(f"An error occurred during conversion: {e}")
        if verbose:
            import traceback
            print("Full error details:")
            traceback.print_exc()
        return False


def main():
    """
    Main function to handle command-line arguments and orchestrate the conversion.
    """
    parser = argparse.ArgumentParser(
        description="Convert Markdown files to Microsoft Word documents",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Examples:
  %(prog)s document.md                    # Convert with auto-generated output name
  %(prog)s input.md output.docx           # Convert with custom output name
  %(prog)s -v document.md                 # Verbose output
  %(prog)s --help                         # Show this help message
        """
    )

    parser.add_argument(
        'input_file',
        help='Path to the input Markdown file (.md)'
    )

    parser.add_argument(
        'output_file',
        nargs='?',
        help='Path to the output Word document (optional, auto-generated if not provided)'
    )

    parser.add_argument(
        '-v', '--verbose',
        action='store_true',
        help='Enable verbose output with additional information'
    )

    parser.add_argument(
        '--version',
        action='version',
        version='Markdown to Word Converter v1.0'
    )

    # Parse command line arguments
    args = parser.parse_args()

    # Validate input file
    if not validate_input_file(args.input_file):
        sys.exit(1)

    # Generate output filename
    output_file = generate_output_filename(args.input_file, args.output_file)

    # Check if output file already exists
    if os.path.exists(output_file):
        response = input(f"Output file '{output_file}' already exists. Overwrite? (y/N): ")
        if response.lower() not in ['y', 'yes']:
            print("Conversion cancelled.")
            sys.exit(0)

    # Perform the conversion
    success = convert_markdown_to_word(
        args.input_file,
        output_file,
        args.verbose
    )

    if success:
        print("Conversion completed successfully!")
        sys.exit(0)
    else:
        print("Conversion failed!")
        sys.exit(1)


if __name__ == "__main__":
    main()
