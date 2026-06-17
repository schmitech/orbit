import io
from typing import Dict, Any

from .base import BaseRenderer


class DocxRenderer(BaseRenderer):
    """Render a document spec to DOCX bytes using python-docx."""

    def render(self, spec: Dict[str, Any]) -> bytes:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        doc = Document()
        title_para = doc.add_heading(spec.get("title", "Document"), level=0)
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        meta_line = self._build_meta_line(spec)
        if meta_line:
            p = doc.add_paragraph()
            run = p.add_run(meta_line)
            run.italic = True
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_paragraph()

        table_style = self._get('docx', 'table_style', default='Table Grid')

        for section in spec.get("sections", []):
            if section.get("heading"):
                doc.add_heading(section["heading"], level=1)
            if section.get("body"):
                doc.add_paragraph(section["body"])
            for point in section.get("bullet_points") or []:
                doc.add_paragraph(point, style="List Bullet")
            table_data = section.get("table")
            if table_data and len(table_data) > 1:
                headers = table_data[0]
                rows = table_data[1:]
                tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
                tbl.style = table_style
                hdr_cells = tbl.rows[0].cells
                for i, h in enumerate(headers):
                    hdr_cells[i].text = str(h)
                    hdr_cells[i].paragraphs[0].runs[0].bold = True
                for r_idx, row in enumerate(rows):
                    for c_idx, val in enumerate(row):
                        tbl.rows[r_idx + 1].cells[c_idx].text = str(val)
                doc.add_paragraph()

        buf = io.BytesIO()
        doc.save(buf)
        return buf.getvalue()
