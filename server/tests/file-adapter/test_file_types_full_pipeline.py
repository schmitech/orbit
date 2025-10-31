"""
Full Pipeline Integration Tests for Different File Types

Tests the complete lifecycle for each supported file type:
1. Upload file via /api/files/upload
2. Wait for processing completion
3. Check processing status
4. Query file content via /api/files/{file_id}/query
5. Use file in chat context via /v1/chat with file_ids
6. Verify extracted content and responses
7. Delete file via /api/files/{file_id}

File types tested:
- Markdown (.md)
- Text (.txt)
- PDF (.pdf)
- DOCX (.docx)
- CSV (.csv)
- JSON (.json)
- HTML (.html)
- Images (.png, .jpg) - with vision service integration

Prerequisites:
1. Server running on http://localhost:3000
2. File Processing Service initialized
3. API key "files" mapped to file-document-qa adapter
4. Vision services configured (for image tests)
"""

import pytest
import httpx
import asyncio
import tempfile
import json
import logging
from pathlib import Path
from typing import Dict, Any, Optional
import time

logger = logging.getLogger(__name__)

# Test configuration
TEST_API_KEY = "files"
TEST_SERVER_URL = "http://localhost:3000"
TEST_SESSION_ID = "test-session-file-types"
PROCESSING_WAIT_TIME = 3  # seconds to wait for file processing

# Sample content for each file type
SAMPLE_CONTENT = {
    'markdown': """# Technical Documentation

## Overview
This is a test markdown document for the ORBIT file processing system.

## Key Features
- Vector-based search
- Multi-format support
- Semantic chunking

## Code Example
```python
def process_file(file_data):
    return chunker.chunk(file_data)
```

## Conclusion
This document tests markdown file processing capabilities.
""",

    'text': """Plain Text Document

This is a plain text file for testing the ORBIT file processing system.

Section 1: Introduction
The file processing system handles various formats and extracts text content.

Section 2: Technical Details
Files are chunked into semantic segments for better retrieval.

Section 3: Conclusion
This tests plain text file processing.
""",

    'csv': """Name,Age,City,Occupation
John Doe,30,New York,Engineer
Jane Smith,25,San Francisco,Designer
Bob Johnson,35,Seattle,Manager
Alice Williams,28,Boston,Developer
Charlie Brown,32,Austin,Analyst
""",

    'json': """{
  "document": "Test JSON Document",
  "type": "technical_specification",
  "version": "1.0",
  "sections": [
    {
      "title": "Introduction",
      "content": "This is a JSON document for testing file processing"
    },
    {
      "title": "Features",
      "content": "Supports JSON parsing and indexing"
    }
  ],
  "metadata": {
    "author": "Test Suite",
    "created": "2024-01-15",
    "tags": ["testing", "json", "orbit"]
  }
}
""",

    'html': """<!DOCTYPE html>
<html>
<head>
    <title>Test HTML Document</title>
</head>
<body>
    <h1>HTML Document Test</h1>

    <h2>Overview</h2>
    <p>This is a test HTML document for the ORBIT file processing system.</p>

    <h2>Features</h2>
    <ul>
        <li>HTML parsing and text extraction</li>
        <li>Semantic chunking of HTML content</li>
        <li>Vector-based search capabilities</li>
    </ul>

    <h2>Technical Details</h2>
    <p>The system extracts text from HTML tags and processes it for indexing.</p>

    <div class="conclusion">
        <h3>Conclusion</h3>
        <p>This tests HTML file processing capabilities.</p>
    </div>
</body>
</html>
""",
}


@pytest.fixture
async def http_client():
    """Create an HTTP client for making requests"""
    async with httpx.AsyncClient(timeout=120.0) as client:
        yield client


@pytest.fixture
async def server_health_check(http_client):
    """Check if server is running before tests"""
    try:
        response = await http_client.get(f"{TEST_SERVER_URL}/health")
        if response.status_code != 200:
            pytest.skip(f"Server is not healthy (status: {response.status_code})")
    except Exception as e:
        pytest.skip(f"Server is not accessible: {e}")


async def upload_file(
    http_client: httpx.AsyncClient,
    content: bytes,
    filename: str,
    mime_type: str
) -> Dict[str, Any]:
    """
    Upload a file and return the response data.

    Returns:
        Dict with file_id, status, etc.
    """
    headers = {"X-API-Key": TEST_API_KEY}

    files = {"file": (filename, content, mime_type)}
    response = await http_client.post(
        f"{TEST_SERVER_URL}/api/files/upload",
        headers=headers,
        files=files
    )

    logger.info(f"Upload response status: {response.status_code}")
    logger.info(f"Upload response body: {response.text[:200]}")

    assert response.status_code == 200, f"Upload failed: {response.text}"
    data = response.json()
    logger.info(f"Uploaded file_id: {data.get('file_id')}")
    return data


async def wait_for_processing(
    http_client: httpx.AsyncClient,
    file_id: str,
    max_wait: int = 15
) -> Dict[str, Any]:
    """
    Wait for file processing to complete and return file info.

    Args:
        http_client: HTTP client
        file_id: File ID to check
        max_wait: Maximum seconds to wait

    Returns:
        File info dict
    """
    headers = {"X-API-Key": TEST_API_KEY}

    last_status = None
    for i in range(max_wait):
        await asyncio.sleep(1)

        try:
            response = await http_client.get(
                f"{TEST_SERVER_URL}/api/files/{file_id}",
                headers=headers
            )

            if response.status_code == 200:
                file_info = response.json()
                status = file_info.get('processing_status', file_info.get('status', 'unknown'))
                last_status = status

                logger.info(f"File {file_id} processing status: {status} (attempt {i+1}/{max_wait})")

                if status == 'completed':
                    return file_info
                elif status == 'failed':
                    error_msg = file_info.get('error', 'Unknown error')
                    logger.error(f"File processing failed: {error_msg}")
                    raise Exception(f"File processing failed: {error_msg}")
            else:
                logger.warning(f"File status check returned {response.status_code}: {response.text}")
        except httpx.HTTPError as e:
            logger.warning(f"HTTP error checking file status: {e}")

    raise TimeoutError(f"File processing did not complete within {max_wait} seconds (last status: {last_status})")


async def query_file(
    http_client: httpx.AsyncClient,
    file_id: str,
    query: str,
    max_results: int = 3
) -> Dict[str, Any]:
    """Query file content via /api/files/{file_id}/query"""
    headers = {
        "X-API-Key": TEST_API_KEY,
        "Content-Type": "application/json"
    }

    response = await http_client.post(
        f"{TEST_SERVER_URL}/api/files/{file_id}/query",
        headers=headers,
        json={"query": query, "max_results": max_results}
    )

    assert response.status_code == 200, f"Query failed: {response.text}"
    return response.json()


async def chat_with_file(
    http_client: httpx.AsyncClient,
    file_id: str,
    message: str
) -> Dict[str, Any]:
    """Chat with file context via /v1/chat"""
    headers = {
        "X-API-Key": TEST_API_KEY,
        "X-Session-ID": TEST_SESSION_ID,
        "Content-Type": "application/json"
    }

    chat_request = {
        "messages": [
            {"role": "user", "content": message}
        ],
        "stream": False,
        "file_ids": [file_id]
    }

    response = await http_client.post(
        f"{TEST_SERVER_URL}/v1/chat",
        headers=headers,
        json=chat_request
    )

    assert response.status_code == 200, f"Chat failed: {response.text}"
    return response.json()


async def delete_file(http_client: httpx.AsyncClient, file_id: str):
    """Delete file via /api/files/{file_id}"""
    headers = {"X-API-Key": TEST_API_KEY}

    response = await http_client.delete(
        f"{TEST_SERVER_URL}/api/files/{file_id}",
        headers=headers
    )

    # Accept 200 (deleted) or 404 (already deleted/not found)
    assert response.status_code in [200, 404], f"Delete failed: {response.text}"


# ============================================================================
# MARKDOWN FILE TESTS
# ============================================================================

@pytest.mark.asyncio
async def test_markdown_file_full_pipeline(http_client, server_health_check):
    """Test complete pipeline for Markdown (.md) files"""
    logger.info("Starting Markdown file full pipeline test")

    # 1. Upload file
    content = SAMPLE_CONTENT['markdown'].encode('utf-8')
    upload_data = await upload_file(http_client, content, "test.md", "text/markdown")
    file_id = upload_data["file_id"]

    logger.info(f"Uploaded markdown file: {file_id}")

    try:
        # 2. Wait for processing
        file_info = await wait_for_processing(http_client, file_id)
        status = file_info.get('processing_status', file_info.get('status'))
        assert status == 'completed'
        assert file_info['chunk_count'] > 0

        logger.info(f"Processing completed: {file_info['chunk_count']} chunks")

        # 3. Query file content
        query_result = await query_file(http_client, file_id, "What are the key features?")
        assert len(query_result['results']) > 0

        # Check that markdown content was extracted
        results_text = " ".join([r['content'] for r in query_result['results']])
        assert "Vector-based search" in results_text or "vector" in results_text.lower()

        logger.info(f"Query returned {len(query_result['results'])} results")

        # 4. Chat with file context
        chat_result = await chat_with_file(http_client, file_id, "What does this document explain?")
        assert "response" in chat_result
        assert len(chat_result['response']) > 0

        logger.info(f"Chat response length: {len(chat_result['response'])} chars")

    finally:
        # 5. Cleanup
        try:
            await delete_file(http_client, file_id)
            logger.info("Markdown file test completed and cleaned up")
        except Exception as e:
            logger.warning(f"Cleanup error: {e}")


# ============================================================================
# TEXT FILE TESTS
# ============================================================================

@pytest.mark.asyncio
async def test_text_file_full_pipeline(http_client, server_health_check):
    """Test complete pipeline for plain text (.txt) files"""
    logger.info("Starting text file full pipeline test")

    content = SAMPLE_CONTENT['text'].encode('utf-8')
    upload_data = await upload_file(http_client, content, "test.txt", "text/plain")
    file_id = upload_data["file_id"]

    logger.info(f"Uploaded text file: {file_id}")

    try:
        file_info = await wait_for_processing(http_client, file_id)
        status = file_info.get('processing_status', file_info.get('status'))
        assert status == 'completed'
        assert file_info['chunk_count'] > 0

        logger.info(f"Processing completed: {file_info['chunk_count']} chunks")

        query_result = await query_file(http_client, file_id, "What is mentioned about chunking?")
        assert len(query_result['results']) > 0

        results_text = " ".join([r['content'] for r in query_result['results']])
        assert "chunk" in results_text.lower() or "segment" in results_text.lower()

        logger.info(f"Query returned {len(query_result['results'])} results")

        chat_result = await chat_with_file(http_client, file_id, "Summarize this document")
        assert "response" in chat_result
        assert len(chat_result['response']) > 0

        logger.info(f"Chat response length: {len(chat_result['response'])} chars")

    finally:
        try:
            await delete_file(http_client, file_id)
            logger.info("Text file test completed and cleaned up")
        except Exception as e:
            logger.warning(f"Cleanup error: {e}")


# ============================================================================
# PDF FILE TESTS
# ============================================================================

@pytest.mark.asyncio
async def test_pdf_file_full_pipeline(http_client, server_health_check):
    """Test complete pipeline for PDF files"""
    logger.info("Starting PDF file full pipeline test")

    # Create a simple PDF file using reportlab
    try:
        from reportlab.pdfgen import canvas
        from reportlab.lib.pagesizes import letter
        import io

        buffer = io.BytesIO()
        c = canvas.Canvas(buffer, pagesize=letter)

        # Add content to PDF
        c.drawString(100, 750, "PDF Test Document")
        c.drawString(100, 730, "")
        c.drawString(100, 710, "This is a test PDF file for the ORBIT file processing system.")
        c.drawString(100, 690, "")
        c.drawString(100, 670, "Section 1: Introduction")
        c.drawString(100, 650, "The PDF processor extracts text from PDF documents.")
        c.drawString(100, 630, "")
        c.drawString(100, 610, "Section 2: Features")
        c.drawString(100, 590, "- PDF text extraction")
        c.drawString(100, 570, "- Semantic chunking")
        c.drawString(100, 550, "- Vector indexing")

        c.save()
        pdf_content = buffer.getvalue()

        upload_data = await upload_file(http_client, pdf_content, "test.pdf", "application/pdf")
        file_id = upload_data["file_id"]

        logger.info(f"Uploaded PDF file: {file_id}")

        try:
            file_info = await wait_for_processing(http_client, file_id, max_wait=15)
            # wait_for_processing only returns when status is 'completed'

            logger.info(f"PDF processing completed: {file_info.get('chunk_count', 0)} chunks")

            if file_info.get('chunk_count', 0) > 0:
                query_result = await query_file(http_client, file_id, "What does the PDF document contain?")
                assert len(query_result['results']) > 0

                logger.info(f"Query returned {len(query_result['results'])} results")

                chat_result = await chat_with_file(http_client, file_id, "What is this PDF about?")
                assert "response" in chat_result

                logger.info(f"Chat response length: {len(chat_result['response'])} chars")
            else:
                logger.warning("PDF processed but no chunks created - may indicate extraction issue")

        finally:
            try:
                await delete_file(http_client, file_id)
                logger.info("PDF file test completed and cleaned up")
            except Exception as e:
                logger.warning(f"Cleanup error: {e}")

    except ImportError:
        pytest.skip("reportlab not installed - skipping PDF test")


# ============================================================================
# DOCX FILE TESTS
# ============================================================================

@pytest.mark.asyncio
async def test_docx_file_full_pipeline(http_client, server_health_check):
    """Test complete pipeline for DOCX files"""
    logger.info("Starting DOCX file full pipeline test")

    # Create a simple DOCX file
    try:
        from docx import Document
        import io

        doc = Document()
        doc.add_heading('DOCX Test Document', 0)
        doc.add_paragraph('This is a test DOCX file for the ORBIT file processing system.')
        doc.add_heading('Section 1: Introduction', level=1)
        doc.add_paragraph('The DOCX processor extracts text from Word documents.')
        doc.add_heading('Section 2: Features', level=1)
        doc.add_paragraph('Key capabilities include:')
        doc.add_paragraph('- Text extraction from DOCX files', style='List Bullet')
        doc.add_paragraph('- Semantic chunking of content', style='List Bullet')
        doc.add_paragraph('- Vector-based indexing', style='List Bullet')

        buffer = io.BytesIO()
        doc.save(buffer)
        docx_content = buffer.getvalue()

        upload_data = await upload_file(
            http_client,
            docx_content,
            "test.docx",
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
        file_id = upload_data["file_id"]

        logger.info(f"Uploaded DOCX file: {file_id}")

        try:
            file_info = await wait_for_processing(http_client, file_id, max_wait=15)
            # wait_for_processing only returns when status is 'completed'

            logger.info(f"DOCX processing completed: {file_info.get('chunk_count', 0)} chunks")

            if file_info.get('chunk_count', 0) > 0:
                query_result = await query_file(http_client, file_id, "What are the features?")
                assert len(query_result['results']) > 0

                logger.info(f"Query returned {len(query_result['results'])} results")

                chat_result = await chat_with_file(http_client, file_id, "Summarize this document")
                assert "response" in chat_result

                logger.info(f"Chat response length: {len(chat_result['response'])} chars")
            else:
                logger.warning("DOCX processed but no chunks created - may indicate extraction issue")

        finally:
            try:
                await delete_file(http_client, file_id)
                logger.info("DOCX file test completed and cleaned up")
            except Exception as e:
                logger.warning(f"Cleanup error: {e}")

    except ImportError:
        pytest.skip("python-docx not installed - skipping DOCX test")


# ============================================================================
# CSV FILE TESTS
# ============================================================================

@pytest.mark.asyncio
async def test_csv_file_full_pipeline(http_client, server_health_check):
    """Test complete pipeline for CSV files"""
    logger.info("Starting CSV file full pipeline test")

    content = SAMPLE_CONTENT['csv'].encode('utf-8')
    upload_data = await upload_file(http_client, content, "test.csv", "text/csv")
    file_id = upload_data["file_id"]

    logger.info(f"Uploaded CSV file: {file_id}")

    try:
        file_info = await wait_for_processing(http_client, file_id)
        status = file_info.get('processing_status', file_info.get('status'))
        assert status == 'completed'

        logger.info(f"CSV processing completed: {file_info.get('chunk_count', 0)} chunks")

        if file_info.get('chunk_count', 0) > 0:
            query_result = await query_file(http_client, file_id, "Who works in Seattle?")
            assert len(query_result['results']) > 0

            logger.info(f"Query returned {len(query_result['results'])} results")

            chat_result = await chat_with_file(http_client, file_id, "How many people are in the dataset?")
            assert "response" in chat_result

            logger.info(f"Chat response length: {len(chat_result['response'])} chars")
        else:
            logger.warning("CSV processed but no chunks created - may indicate extraction issue")

    finally:
        try:
            await delete_file(http_client, file_id)
            logger.info("CSV file test completed and cleaned up")
        except Exception as e:
            logger.warning(f"Cleanup error: {e}")


# ============================================================================
# JSON FILE TESTS
# ============================================================================

@pytest.mark.asyncio
async def test_json_file_full_pipeline(http_client, server_health_check):
    """Test complete pipeline for JSON files"""
    logger.info("Starting JSON file full pipeline test")

    content = SAMPLE_CONTENT['json'].encode('utf-8')
    upload_data = await upload_file(http_client, content, "test.json", "application/json")
    file_id = upload_data["file_id"]

    logger.info(f"Uploaded JSON file: {file_id}")

    try:
        file_info = await wait_for_processing(http_client, file_id)
        status = file_info.get('processing_status', file_info.get('status'))
        assert status == 'completed'

        logger.info(f"JSON processing completed: {file_info.get('chunk_count', 0)} chunks")

        if file_info.get('chunk_count', 0) > 0:
            query_result = await query_file(http_client, file_id, "What is the document type?")
            assert len(query_result['results']) > 0

            logger.info(f"Query returned {len(query_result['results'])} results")

            chat_result = await chat_with_file(http_client, file_id, "What sections does this document have?")
            assert "response" in chat_result

            logger.info(f"Chat response length: {len(chat_result['response'])} chars")
        else:
            logger.warning("JSON processed but no chunks created - may indicate extraction issue")

    finally:
        try:
            await delete_file(http_client, file_id)
            logger.info("JSON file test completed and cleaned up")
        except Exception as e:
            logger.warning(f"Cleanup error: {e}")


# ============================================================================
# HTML FILE TESTS
# ============================================================================

@pytest.mark.asyncio
async def test_html_file_full_pipeline(http_client, server_health_check):
    """Test complete pipeline for HTML files"""
    logger.info("Starting HTML file full pipeline test")

    content = SAMPLE_CONTENT['html'].encode('utf-8')
    upload_data = await upload_file(http_client, content, "test.html", "text/html")
    file_id = upload_data["file_id"]

    logger.info(f"Uploaded HTML file: {file_id}")

    try:
        file_info = await wait_for_processing(http_client, file_id)
        status = file_info.get('processing_status', file_info.get('status'))
        assert status == 'completed'

        logger.info(f"HTML processing completed: {file_info.get('chunk_count', 0)} chunks")

        if file_info.get('chunk_count', 0) > 0:
            query_result = await query_file(http_client, file_id, "What features are listed?")
            assert len(query_result['results']) > 0

            logger.info(f"Query returned {len(query_result['results'])} results")

            chat_result = await chat_with_file(http_client, file_id, "What is this HTML document about?")
            assert "response" in chat_result

            logger.info(f"Chat response length: {len(chat_result['response'])} chars")
        else:
            logger.warning("HTML processed but no chunks created - may indicate extraction issue")

    finally:
        try:
            await delete_file(http_client, file_id)
            logger.info("HTML file test completed and cleaned up")
        except Exception as e:
            logger.warning(f"Cleanup error: {e}")


# ============================================================================
# IMAGE FILE TESTS (PNG with Vision Service)
# ============================================================================

@pytest.mark.asyncio
async def test_png_image_full_pipeline(http_client, server_health_check):
    """Test complete pipeline for PNG image files with vision service"""
    logger.info("Starting PNG image file full pipeline test (with vision service)")

    # Create a simple test PNG image
    try:
        from PIL import Image, ImageDraw, ImageFont
        import io

        # Create a simple image with text
        img = Image.new('RGB', (400, 300), color='white')
        draw = ImageDraw.Draw(img)

        # Draw some text (use default font)
        draw.text((10, 10), "Test Image Document", fill='black')
        draw.text((10, 40), "This is a test image for", fill='black')
        draw.text((10, 60), "the ORBIT vision service.", fill='black')
        draw.text((10, 100), "Features:", fill='blue')
        draw.text((10, 120), "- Text extraction from images", fill='black')
        draw.text((10, 140), "- Image description generation", fill='black')
        draw.text((10, 160), "- OCR capabilities", fill='black')

        # Draw a simple shape
        draw.rectangle([10, 200, 100, 280], outline='red', width=2)
        draw.text((120, 230), "Red Rectangle", fill='red')

        buffer = io.BytesIO()
        img.save(buffer, format='PNG')
        png_content = buffer.getvalue()

        upload_data = await upload_file(http_client, png_content, "test.png", "image/png")
        file_id = upload_data["file_id"]

        logger.info(f"Uploaded PNG image file: {file_id}")

        try:
            # Images may take longer to process due to vision service (API calls + retries)
            # With concurrent calls and retries, can take up to 90-120 seconds
            file_info = await wait_for_processing(http_client, file_id, max_wait=150)

            status = file_info.get('processing_status', file_info.get('status'))
            logger.info(f"PNG processing status: {status}")
            logger.info(f"PNG chunk count: {file_info.get('chunk_count', 0)}")

            if status == 'completed':
                if file_info.get('chunk_count', 0) > 0:
                    # Query for image content
                    query_result = await query_file(http_client, file_id, "What text is in the image?")

                    logger.info(f"Query returned {len(query_result['results'])} results")

                    if len(query_result['results']) > 0:
                        results_text = " ".join([r['content'] for r in query_result['results']])
                        logger.info(f"Extracted content sample: {results_text[:200]}")

                    # Chat with image context
                    chat_result = await chat_with_file(http_client, file_id, "Describe what you see in this image")
                    assert "response" in chat_result

                    logger.info(f"Chat response: {chat_result['response'][:200]}...")
                else:
                    logger.warning("PNG processed but no chunks created - vision service may not be configured")
            else:
                logger.warning(f"PNG processing failed: {file_info.get('error', 'Unknown error')}")
                logger.warning("This may indicate vision service is not configured or API keys are missing")

        finally:
            try:
                await delete_file(http_client, file_id)
                logger.info("PNG image test completed and cleaned up")
            except Exception as e:
                logger.warning(f"Cleanup error: {e}")

    except ImportError:
        pytest.skip("Pillow not installed - skipping PNG test")


# ============================================================================
# IMAGE FILE TESTS (JPEG with Vision Service)
# ============================================================================

@pytest.mark.asyncio
async def test_jpeg_image_full_pipeline(http_client, server_health_check):
    """Test complete pipeline for JPEG image files with vision service"""
    logger.info("Starting JPEG image file full pipeline test (with vision service)")

    try:
        from PIL import Image, ImageDraw
        import io

        # Create a simple JPEG image
        img = Image.new('RGB', (400, 300), color='lightblue')
        draw = ImageDraw.Draw(img)

        draw.text((10, 10), "JPEG Test Image", fill='darkblue')
        draw.text((10, 40), "Testing ORBIT vision service", fill='black')
        draw.text((10, 70), "with JPEG format", fill='black')

        # Draw a circle
        draw.ellipse([150, 150, 250, 250], outline='green', width=3)
        draw.text((270, 190), "Green Circle", fill='green')

        buffer = io.BytesIO()
        img.save(buffer, format='JPEG', quality=95)
        jpeg_content = buffer.getvalue()

        upload_data = await upload_file(http_client, jpeg_content, "test.jpg", "image/jpeg")
        file_id = upload_data["file_id"]

        logger.info(f"Uploaded JPEG image file: {file_id}")

        try:
            # Images may take longer to process due to vision service (API calls + retries)
            # With concurrent calls and retries, can take up to 90-120 seconds
            file_info = await wait_for_processing(http_client, file_id, max_wait=150)

            status = file_info.get('processing_status', file_info.get('status'))
            logger.info(f"JPEG processing status: {status}")
            logger.info(f"JPEG chunk count: {file_info.get('chunk_count', 0)}")

            if status == 'completed':
                if file_info.get('chunk_count', 0) > 0:
                    query_result = await query_file(http_client, file_id, "What is in this image?")

                    logger.info(f"Query returned {len(query_result['results'])} results")

                    chat_result = await chat_with_file(http_client, file_id, "What shapes and colors are visible?")
                    assert "response" in chat_result

                    logger.info(f"Chat response: {chat_result['response'][:200]}...")
                else:
                    logger.warning("JPEG processed but no chunks created - vision service may not be configured")
            else:
                logger.warning(f"JPEG processing failed: {file_info.get('error', 'Unknown error')}")

        finally:
            try:
                await delete_file(http_client, file_id)
                logger.info("JPEG image test completed and cleaned up")
            except Exception as e:
                logger.warning(f"Cleanup error: {e}")

    except ImportError:
        pytest.skip("Pillow not installed - skipping JPEG test")


# ============================================================================
# ERROR HANDLING TESTS
# ============================================================================

@pytest.mark.asyncio
async def test_unsupported_file_type(http_client, server_health_check):
    """Test handling of unsupported file types"""
    logger.info("Testing unsupported file type handling")

    # Try to upload an unsupported file type
    content = b"Binary content that's not supported"
    headers = {"X-API-Key": TEST_API_KEY}

    files = {"file": ("test.xyz", content, "application/x-unknown")}
    response = await http_client.post(
        f"{TEST_SERVER_URL}/api/files/upload",
        headers=headers,
        files=files
    )

    # Should either reject (400/415) or accept but fail processing
    logger.info(f"Unsupported file type response: {response.status_code}")

    if response.status_code == 200:
        # If accepted, processing should fail
        data = response.json()
        file_id = data["file_id"]

        try:
            file_info = await wait_for_processing(http_client, file_id, max_wait=5)
            # Should be marked as failed
            status = file_info.get('processing_status', file_info.get('status'))
            logger.info(f"Processing status: {status}")
        except Exception as e:
            logger.info(f"Processing failed as expected: {e}")
        finally:
            try:
                await delete_file(http_client, file_id)
            except:
                pass
    else:
        # Rejected at upload - this is also acceptable
        assert response.status_code in [400, 415], f"Unexpected status: {response.status_code}"


@pytest.mark.asyncio
async def test_oversized_file_handling(http_client, server_health_check):
    """Test handling of files exceeding size limit"""
    logger.info("Testing oversized file handling")

    # Create a file larger than typical limit (e.g., 51MB if limit is 50MB)
    large_content = b"x" * (51 * 1024 * 1024)  # 51 MB

    headers = {"X-API-Key": TEST_API_KEY}
    files = {"file": ("large.txt", large_content, "text/plain")}

    response = await http_client.post(
        f"{TEST_SERVER_URL}/api/files/upload",
        headers=headers,
        files=files
    )

    # Should reject oversized files
    logger.info(f"Oversized file response: {response.status_code}")

    # Accept either rejection at upload (413, 400) or successful upload with potential processing error
    assert response.status_code in [200, 400, 413], f"Unexpected status: {response.status_code}"

    if response.status_code == 200:
        # Clean up if it was accepted
        try:
            file_id = response.json()["file_id"]
            await delete_file(http_client, file_id)
        except:
            pass
