"""
DOCX Processor

Handles DOCX files using python-docx.
"""

import logging
from typing import Dict, Any
from .base_processor import FileProcessor

logger = logging.getLogger(__name__)

try:
    from docx import Document
    from io import BytesIO
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    logger.warning("python-docx not available. DOCX processing disabled.")


class DOCXProcessor(FileProcessor):
    """
    Processor for DOCX files.
    
    Supports: application/vnd.openxmlformats-officedocument.wordprocessingml.document
    Requires: python-docx
    """
    
    def supports_mime_type(self, mime_type: str) -> bool:
        """Check if this processor supports the MIME type."""
        docx_types = [
            'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            'application/msword',
        ]
        return DOCX_AVAILABLE and mime_type.lower() in docx_types
    
    async def extract_text(self, file_data: bytes, filename: str = None) -> str:
        """Extract text from DOCX."""
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx not available")
        
        text_parts = []
        
        try:
            doc = Document(BytesIO(file_data))
            
            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text_parts.append(cell.text)
            
            return "\n\n".join(text_parts)
        
        except Exception as e:
            self.logger.error(f"Error processing DOCX: {e}")
            raise
    
    async def extract_metadata(self, file_data: bytes, filename: str = None) -> Dict[str, Any]:
        """Extract metadata from DOCX."""
        metadata = await super().extract_metadata(file_data, filename)
        
        if not DOCX_AVAILABLE:
            return metadata
        
        try:
            doc = Document(BytesIO(file_data))
            
            # Count paragraphs and tables
            paragraph_count = len(doc.paragraphs)
            table_count = len(doc.tables)
            
            metadata.update({
                'paragraph_count': paragraph_count,
                'table_count': table_count,
                'mime_type': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            })
        
        except Exception as e:
            self.logger.warning(f"Error extracting DOCX metadata: {e}")
        
        return metadata
